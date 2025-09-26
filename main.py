from fastapi import FastAPI
from fastapi.responses import Response
from pydantic import BaseModel
from io import BytesIO
from zipfile import ZipFile
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from docx import Document
from docx.shared import Inches

app = FastAPI()

PAGE_WIDTH, PAGE_HEIGHT = A4
TABLE_WIDTH = 170 * mm

class ProposalData(BaseModel):
    website: str
    company_name: str
    email: str
    organization: str
    address: str
    gst_number: str
    contract_no: str
    date: str
    seller: str
    consignee: str
    notify_party: str
    product: str
    quantity: str
    price: str
    amount: str
    packing: str
    loading_port: str
    destination_port: str
    shipment: str
    sellers_bank: str
    account_no: str
    documents: str
    payment_terms: str
    arbitration: str
    terms: list[str]

def generate_pdf(data: ProposalData) -> BytesIO:
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            topMargin=20*mm, bottomMargin=20*mm,
                            leftMargin=18*mm, rightMargin=18*mm)

    styles = getSampleStyleSheet()
    s_title = ParagraphStyle("Title", parent=styles["Heading1"], alignment=TA_CENTER, fontSize=16, spaceAfter=8)
    s_h1 = ParagraphStyle("H1", parent=styles["Heading2"], alignment=TA_LEFT, fontSize=12, spaceAfter=6, leading=14)
    s_normal = ParagraphStyle("Normal", parent=styles["Normal"], fontSize=10, leading=13)
    s_bold = ParagraphStyle("Bold", parent=styles["Normal"], fontSize=10, leading=13, fontName="Helvetica-Bold")
    s_product = ParagraphStyle("Product", parent=styles["Normal"], fontSize=9, leading=11)

    elements = []

    # Header
    header_data = [
        [Paragraph(f"<b>Website:</b> {data.website}", s_normal),
         Paragraph(f"<b>{data.company_name.upper()}</b>", s_title),
         Paragraph(f"<b>Email:</b> {data.email}", s_normal)]
    ]
    header_table = Table(header_data, colWidths=[TABLE_WIDTH/3]*3, hAlign="CENTER")
    header_table.setStyle(TableStyle([("VALIGN", (0,0), (-1,-1), "MIDDLE")]))
    elements.append(header_table)
    elements.append(Spacer(1, 6))

    # Organization & Address/GST
    elements.append(Paragraph(data.organization, s_normal))
    elements.append(Spacer(1, 10))
    addr_data = [
        [Paragraph(f"<b>Address:</b> {data.address}", s_normal),
         Paragraph(f"<b>GST:</b> {data.gst_number}", s_normal)]
    ]
    addr_table = Table(addr_data, colWidths=[TABLE_WIDTH*0.65, TABLE_WIDTH*0.35], hAlign="CENTER")
    elements.append(addr_table)
    elements.append(Spacer(1, 12))

    # Title & Contract
    elements.append(Paragraph("<b>SALES CONTRACT PROPOSAL</b>", s_title))
    contract_data = [
        [Paragraph(f"<b>Contract No:</b> {data.contract_no}", s_normal),
         Paragraph(f"<b>Date:</b> {data.date}", s_normal)]
    ]
    contract_table = Table(contract_data, colWidths=[TABLE_WIDTH/2, TABLE_WIDTH/2], hAlign="CENTER")
    elements.append(contract_table)
    elements.append(Spacer(1, 16))

    # Parties
    parties_data = [
        [Paragraph(f"<b>SELLER</b><br/>{data.seller}", s_normal),
         Paragraph(f"<b>CONSIGNEE</b><br/>{data.consignee}", s_normal),
         Paragraph(f"<b>NOTIFY PARTY</b><br/>{data.notify_party}", s_normal)]
    ]
    parties_table = Table(parties_data, colWidths=[TABLE_WIDTH/3]*3, hAlign="CENTER")
    parties_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.black),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
    ]))
    elements.append(parties_table)
    elements.append(Spacer(1, 16))

    # Product Table
    product_data = [
        [Paragraph("<b>Product</b>", s_bold),
         Paragraph("<b>Quantity</b>", s_bold),
         Paragraph("<b>Price (CIF)</b>", s_bold),
         Paragraph("<b>Amount (CIF)</b>", s_bold)],
        [Paragraph(data.product, s_product),
         Paragraph(data.quantity, s_product),
         Paragraph(data.price, s_product),
         Paragraph(data.amount, s_product)]
    ]
    prod_table = Table(product_data, colWidths=[TABLE_WIDTH*0.35, TABLE_WIDTH*0.15, TABLE_WIDTH*0.25, TABLE_WIDTH*0.25], hAlign="CENTER")
    prod_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.5, colors.grey),
        ("BACKGROUND", (0,0), (-1,0), colors.lightblue),
        ("VALIGN", (0,0), (-1,-1), "MIDDLE"),
    ]))
    elements.append(prod_table)
    elements.append(Spacer(1, 15))

    # Shipment & Payment
    shipment_data = [
        ["Packing", data.packing],
        ["Loading Port", data.loading_port],
        ["Destination Port", data.destination_port],
        ["Shipment", data.shipment],
        ["Seller’s Bank", data.sellers_bank],
        ["Account No.", data.account_no],
        ["Documents", data.documents],
        ["Payment Terms", data.payment_terms],
    ]
    shipment_table = Table(shipment_data, colWidths=[TABLE_WIDTH/3, (TABLE_WIDTH*2)/3], hAlign="CENTER")
    shipment_table.setStyle(TableStyle([
        ("GRID", (0,0), (-1,-1), 0.4, colors.lightgrey),
        ("BACKGROUND", (0,0), (0,-1), colors.whitesmoke),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
    ]))
    elements.append(shipment_table)
    elements.append(Spacer(1, 16))

    # Arbitration
    elements.append(Paragraph("<b>Arbitration</b>", s_h1))
    elements.append(Paragraph(data.arbitration, s_normal))
    elements.append(Spacer(1, 12))

    # Terms
    elements.append(Paragraph("<b>Terms & Conditions</b>", s_h1))
    for t in data.terms:
        elements.append(Paragraph(t, s_normal))
    elements.append(Spacer(1, 24))

    # Signatures
    sign_data = [
        [Paragraph("<b>Accepted</b>", s_bold), "", ""],
        [Paragraph("For, Seller", s_normal), Paragraph("For, Consignee", s_normal), Paragraph("For, Notify Party", s_normal)],
        ["\n\n\n____________________", "\n\n\n____________________", "\n\n\n____________________"],
    ]
    sign_table = Table(sign_data, colWidths=[TABLE_WIDTH/3]*3, hAlign="CENTER")
    elements.append(sign_table)

    doc.build(elements)
    buffer.seek(0)
    return buffer

def generate_word(data: ProposalData) -> BytesIO:
    doc = Document()
    doc.add_heading(data.company_name.upper(), 0)
    doc.add_paragraph(f"Website: {data.website} | Email: {data.email}")
    doc.add_paragraph(data.organization)
    doc.add_paragraph(f"Address: {data.address}")
    doc.add_paragraph(f"GST: {data.gst_number}")
    doc.add_paragraph(f"Contract No: {data.contract_no} | Date: {data.date}")

    doc.add_heading("SELLER", level=1)
    doc.add_paragraph(data.seller)
    doc.add_heading("CONSIGNEE", level=1)
    doc.add_paragraph(data.consignee)
    doc.add_heading("NOTIFY PARTY", level=1)
    doc.add_paragraph(data.notify_party)

    table = doc.add_table(rows=2, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Product"
    hdr_cells[1].text = "Quantity"
    hdr_cells[2].text = "Price (CIF)"
    hdr_cells[3].text = "Amount (CIF)"
    row_cells = table.rows[1].cells
    row_cells[0].text = data.product
    row_cells[1].text = data.quantity
    row_cells[2].text = data.price
    row_cells[3].text = data.amount

    doc.add_paragraph(f"Packing: {data.packing}")
    doc.add_paragraph(f"Loading Port: {data.loading_port}")
    doc.add_paragraph(f"Destination Port: {data.destination_port}")
    doc.add_paragraph(f"Shipment: {data.shipment}")
    doc.add_paragraph(f"Seller’s Bank: {data.sellers_bank}")
    doc.add_paragraph(f"Account No.: {data.account_no}")
    doc.add_paragraph(f"Documents: {data.documents}")
    doc.add_paragraph(f"Payment Terms: {data.payment_terms}")
    doc.add_paragraph("Arbitration: " + data.arbitration)
    doc.add_heading("Terms & Conditions", level=1)
    for t in data.terms:
        doc.add_paragraph(t)
    doc.add_paragraph("Accepted\nFor, Seller\nFor, Consignee\nFor, Notify Party")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.post("/generate-proposal/")
async def generate_proposal(data: ProposalData):
    pdf_buffer = generate_pdf(data)
    word_buffer = generate_word(data)

    zip_buffer = BytesIO()
    with ZipFile(zip_buffer, "w") as zip_file:
        zip_file.writestr("Proposal.pdf", pdf_buffer.read())
        zip_file.writestr("Proposal.docx", word_buffer.read())
    zip_buffer.seek(0)

    return Response(
        content=zip_buffer.read(),
        media_type="application/zip",
        headers={"Content-Disposition": "attachment; filename=Proposal.zip"}
    )
